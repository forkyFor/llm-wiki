"""OCR service — multi-backend text extraction.

Backend selection via OCR_BACKEND env var:

  tesseract (default)
    Tier 1 — PDF embedded text via pypdfium2 (instant, lossless)
    Tier 2 — Scanned PDF/image → pytesseract (ita+eng)
    Tier 3 — DOCX/DOC via python-docx
    Tier 4 — Plain text UTF-8

  easyocr
    Tier 1 — PDF embedded text via pypdfium2
    Tier 2 — Scanned PDF/image → EasyOCR (ita+en) — better accuracy than Tesseract
              Fallback to Tesseract if EasyOCR not installed / fails
    Tier 3 — DOCX/DOC
    Tier 4 — Plain text UTF-8

  got_ocr2
    Tier 1 — PDF embedded text via pypdfium2
    Tier 2 — Scanned PDF/image → GOT-OCR2 (stepfun-ai/GOT-OCR-2.0-hf, ~580 MB)
              Full-page OCR — handles complex layouts, multi-column, tables
              Multilingual. First use: auto-downloads ~580 MB.
              Fallback to Tesseract if transformers/torch not installed or fails
    Tier 3 — DOCX/DOC
    Tier 4 — Plain text UTF-8

  ollama
    Tier 1 — PDF embedded text via pypdfium2
    Tier 2 — Scanned PDF/image → Ollama vision model (OCR_MODEL_NAME, default minicpm-v)
              Best quality, slowest — handles complex layouts, tables, handwriting
              Fallback to Tesseract if Ollama vision fails
    Tier 3 — DOCX/DOC
    Tier 4 — Plain text UTF-8

Install:
  tesseract:  winget install UB-Mannheim.TesseractOCR  (Windows)
              apt install tesseract-ocr tesseract-ocr-ita  (Linux)
  easyocr:    pip install easyocr  (~120 MB model auto-download on first use)
  got_ocr2:   pip install transformers torch torchvision  (~580 MB model auto-download on first use)
  ollama:     ollama pull minicpm-v  (5.5 GB)
"""
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

_EMBEDDED_TEXT_MIN_CHARS = 50

# Lazy singletons
_tesseract_available: bool | None = None
_easyocr_reader = None


# ── Tesseract ──────────────────────────────────────────────────────────────

def _check_tesseract() -> bool:
    global _tesseract_available
    if _tesseract_available is not None:
        return _tesseract_available
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        _tesseract_available = True
        logger.info("Tesseract disponibile: %s", pytesseract.get_tesseract_version())
    except Exception as exc:
        logger.warning(
            "Tesseract non disponibile — installa con: winget install UB-Mannheim.TesseractOCR  |  %s", exc
        )
        _tesseract_available = False
    return _tesseract_available


def _ocr_tesseract(image_bytes: bytes) -> str:
    if not _check_tesseract():
        return ""
    import pytesseract
    from PIL import Image
    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    return pytesseract.image_to_string(image, lang="ita+eng")


# ── EasyOCR ────────────────────────────────────────────────────────────────

def _get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        logger.info("Loading EasyOCR (first run: downloading models ~120 MB)…")
        _easyocr_reader = easyocr.Reader(["it", "en"], gpu=False, verbose=False)
        logger.info("EasyOCR ready")
    return _easyocr_reader


def _ocr_easyocr(image_bytes: bytes) -> str:
    try:
        import numpy as np
        from PIL import Image
        reader = _get_easyocr_reader()
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        results = reader.readtext(image, detail=0, paragraph=True)
        return "\n".join(results)
    except ImportError:
        logger.warning("EasyOCR non installato — fallback Tesseract. pip install easyocr")
        return _ocr_tesseract(image_bytes)
    except Exception as exc:
        logger.warning("EasyOCR fallito: %s — fallback Tesseract", exc)
        return _ocr_tesseract(image_bytes)


# ── GOT-OCR2 ──────────────────────────────────────────────────────────────

_got_ocr2_model = None
_got_ocr2_tokenizer = None


def _get_got_ocr2():
    global _got_ocr2_model, _got_ocr2_tokenizer
    if _got_ocr2_model is None:
        from transformers import AutoTokenizer, AutoModelForCausalLM
        import torch
        model_id = "stepfun-ai/GOT-OCR-2.0-hf"
        logger.info("Loading GOT-OCR2 (first run: downloading ~580 MB)…")
        # trust_remote_code required by GOT-OCR-2.0-hf architecture  # nosec B615
        _got_ocr2_tokenizer = AutoTokenizer.from_pretrained(  # nosec B615
            model_id, trust_remote_code=True, revision="main"
        )
        _got_ocr2_model = AutoModelForCausalLM.from_pretrained(  # nosec B615
            model_id,
            trust_remote_code=True,
            revision="main",
            torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
            device_map="auto",
        )
        _got_ocr2_model.eval()
        logger.info("GOT-OCR2 ready")
    return _got_ocr2_model, _got_ocr2_tokenizer


def _ocr_got_ocr2(image_bytes: bytes) -> str:
    try:
        import tempfile
        import os
        from PIL import Image

        model, tokenizer = _get_got_ocr2()
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
            image.save(tmp_path, format="PNG")

        try:
            result = model.chat(tokenizer, tmp_path, ocr_type="ocr")
        finally:
            os.unlink(tmp_path)

        return result if isinstance(result, str) else str(result)
    except ImportError:
        logger.warning("GOT-OCR2: transformers/torch non installati — fallback Tesseract. pip install transformers torch torchvision")
        return _ocr_tesseract(image_bytes)
    except Exception as exc:
        logger.warning("GOT-OCR2 fallito: %s — fallback Tesseract", exc)
        return _ocr_tesseract(image_bytes)


# ── Ollama vision ──────────────────────────────────────────────────────────

def _ocr_ollama_vision(image_bytes: bytes) -> str:
    """Call Ollama vision model to extract text from image (sync, runs in executor)."""
    try:
        import base64
        import httpx
        from backend.config import settings

        b64 = base64.b64encode(image_bytes).decode()
        payload = {
            "model": settings.ocr_model_name,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Extract ALL text visible in this image exactly as it appears. "
                                "Preserve layout, line breaks, and formatting. "
                                "Return only the extracted text — no commentary, no descriptions."
                            ),
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{b64}"},
                        },
                    ],
                }
            ],
            "stream": False,
            "options": {"temperature": 0},
        }
        with httpx.Client(timeout=600) as client:
            resp = client.post(f"{settings.ollama_url}/v1/chat/completions", json=payload)
            resp.raise_for_status()
            text = resp.json()["choices"][0]["message"]["content"]
            logger.debug("Ollama vision OCR: %d chars estratti", len(text))
            return text
    except Exception as exc:
        logger.warning("Ollama vision OCR fallito: %s — fallback Tesseract", exc)
        return _ocr_tesseract(image_bytes)


# ── Backend dispatcher ─────────────────────────────────────────────────────

def _ocr_image(image_bytes: bytes) -> str:
    """OCR image bytes using the configured backend."""
    from backend.config import settings
    backend = settings.ocr_backend

    if backend == "easyocr":
        return _ocr_easyocr(image_bytes)
    if backend == "got_ocr2":
        return _ocr_got_ocr2(image_bytes)
    if backend == "ollama":
        return _ocr_ollama_vision(image_bytes)
    return _ocr_tesseract(image_bytes)  # default: tesseract


# ── PDF helpers ────────────────────────────────────────────────────────────

def _extract_pdf_embedded_text(pdf_path: Path) -> str:
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(str(pdf_path))
    parts: list[str] = []
    for page in pdf:
        textpage = page.get_textpage()
        parts.append(textpage.get_text_range())
    return "\n\n".join(parts)


def _pdf_to_images(pdf_path: Path) -> list[bytes]:
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(str(pdf_path))
    images: list[bytes] = []
    for page in pdf:
        bitmap = page.render(scale=2)
        pil_image = bitmap.to_pil()
        buf = io.BytesIO()
        pil_image.save(buf, format="PNG")
        images.append(buf.getvalue())
    return images


def _extract_pdf(pdf_path: Path) -> str:
    # Tier 1 — embedded text (fast, lossless)
    try:
        embedded = _extract_pdf_embedded_text(pdf_path)
        if len(embedded.strip()) >= _EMBEDDED_TEXT_MIN_CHARS:
            logger.debug("PDF %s: testo embedded OK (%d chars)", pdf_path.name, len(embedded))
            return embedded
        logger.info(
            "PDF %s: testo embedded scarso (%d chars) — provo OCR backend",
            pdf_path.name, len(embedded.strip()),
        )
    except Exception as exc:
        logger.warning("PDF %s: estrazione testo fallita: %s — provo OCR backend", pdf_path.name, exc)

    # Tier 2 — rasterize → OCR backend
    try:
        page_images = _pdf_to_images(pdf_path)
    except Exception as exc:
        logger.error("PDF %s: rasterizzazione fallita: %s", pdf_path.name, exc)
        return ""

    from backend.config import settings
    from concurrent.futures import ThreadPoolExecutor

    n = len(page_images)
    logger.info("PDF %s: OCR backend=%s su %d pagine (parallelo)", pdf_path.name, settings.ocr_backend, n)

    def _safe_ocr(img_bytes: bytes) -> str:
        try:
            return _ocr_image(img_bytes)
        except Exception as exc:
            logger.warning("PDF %s: pagina OCR fallita: %s", pdf_path.name, exc)
            return ""

    max_workers = min(4, n)
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        parts = list(pool.map(_safe_ocr, page_images))

    return "\n\n".join(parts)


# ── DOCX ───────────────────────────────────────────────────────────────────

def _extract_docx(docx_path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(docx_path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        logger.warning("python-docx non installato — .docx non supportato. pip install python-docx")
        return ""
    except Exception as exc:
        logger.error("Errore lettura .docx %s: %s", docx_path.name, exc)
        return ""


# ── Public API ─────────────────────────────────────────────────────────────

def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        return _extract_pdf(file_path)

    if suffix in {".docx", ".doc"}:
        return _extract_docx(file_path)

    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"}:
        return _ocr_image(file_path.read_bytes())

    try:
        return _ocr_image(file_path.read_bytes())
    except Exception:
        logger.warning("Tipo file non supportato: %s", suffix)
        return ""
